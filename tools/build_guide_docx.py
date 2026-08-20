#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate editable Word (.docx) versions of the study guides and cram sheets.

Jaxon, 2026-08-20: "add to all the guides Word doc downloadable versions? make
sure they look nice after downloading and they keep the structure but i want it
to be able to be edited if people wanna change stuff. only for semester 2 and on"

So three requirements, in order of how easy they are to get wrong:

  EDITABLE   Everything uses real Word styles -- Heading 1/2/3, Normal, a named
             table style -- rather than hand-applied formatting. A classmate who
             wants everything smaller changes the style once and the whole
             document follows. Nothing is an image of text.
  STRUCTURED The heading hierarchy is preserved, so Word's navigation pane and
             the TOC field at the top both work. Tables stay tables, lists stay
             lists, figures keep their captions and slide citations.
  NICE       Per-guide accent colour lifted from the page itself, shaded table
             headers, the instructional-objective and professor-emphasis boxes
             rendered as real shaded boxes rather than being flattened away.

SCOPE is read from semesters.js, not hardcoded: whichever classes belong to
fall-2026 or later get exports. Summer 1 is deliberately excluded.

WHY PRE-GENERATED rather than converted in the browser: fidelity is guaranteed,
there is no library to vendor into a public static site, the file works offline,
and the download is a plain <a download> that every browser handles.

    python3 tools/build_guide_docx.py            # every in-scope document
    python3 tools/build_guide_docx.py --list     # show scope and stop
"""
import os, re, sys, io as _io, json, urllib.parse

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)

try:
    from bs4 import BeautifulSoup, NavigableString, Tag
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from PIL import Image
except ImportError as e:                                  # fail loudly, never silently
    sys.exit("missing dependency: %s\n  pip install python-docx beautifulsoup4 pillow" % e)

# ---------------------------------------------------------------- scope

def semester_scope():
    """Classes in fall-2026 or later, read from semesters.js.

    Hardcoding the list would drift the moment a term is added, which is the
    exact failure semesters.js exists to prevent.
    """
    s = open(os.path.join(ROOT, "semesters.js"), encoding="utf-8").read()
    ids = re.findall(r'id:\s*"([^"]+)"', s)
    order = {sid: i for i, sid in enumerate(ids)}
    cutoff = order.get("fall-2026")
    assert cutoff is not None, "fall-2026 not found in semesters.js"
    out = set()
    for m in re.finditer(r'id:\s*"([^"]+)"(.*?)classes:\s*\[(.*?)\]', s, re.S):
        sid, _, cls = m.groups()
        if order[sid] >= cutoff:
            out.update(re.findall(r'"([^"]+)"', cls))
    return out


# class slug -> the repo folder prefix its documents live under
CLASS_DIRS = {
    "cms-1":                "Clinical Medicine and Surgery I Exam",
    "pdm-1":                "Principles of Diagnostic Medicine I Exam",
    "microbiology":         "Microbiology Exam",
    "pharm-1":              "Pharmacology I Exam",
    "physical-diagnosis-2": "Physical Diagnosis 2 Exam",
    "clin-path-1":          "Clinical Pathophysiology I Exam",
    "med-lit":              "Interpretation of Medical Literature Exam",
}


def in_scope_documents():
    """Every document CARDED ON guides.html that belongs to a Semester 2+ class.

    Jaxon, 2026-08-20: "I wanted for all the guides pages even the derm chart."
    So the list is taken from guides.html itself rather than from filename
    patterns -- that page IS the definition of "the guides pages", and driving
    off it means a new reference sheet is picked up the moment it is carded,
    with no pattern here to remember to extend. The earlier version matched
    *study-guide* / *cram-sheet* and silently skipped the comparison chart and
    the three OSCE sheets.
    """
    slugs = semester_scope()
    prefixes = tuple(CLASS_DIRS[s] for s in slugs if s in CLASS_DIRS)
    page = open(os.path.join(ROOT, "guides.html"), encoding="utf-8").read()
    seen, docs = set(), []
    for m in re.finditer(r'class="guide-card[^"]*"\s+href="([^"]+)"', page):
        href = urllib.parse.unquote(m.group(1))
        if not href.endswith(".html") or not href.startswith(prefixes):
            continue
        if href in seen or not os.path.exists(os.path.join(ROOT, href)):
            continue
        seen.add(href)
        docs.append((href.split("/")[0], href))
    return sorted(docs, key=lambda t: t[1])


# ---------------------------------------------------------------- helpers

def shade(cell_or_para, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear"); el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hexcolor)
    (cell_or_para._tc if hasattr(cell_or_para, "_tc") else cell_or_para._p).\
        get_or_add_tcPr() if False else None
    target = cell_or_para._tc.get_or_add_tcPr() if hasattr(cell_or_para, "_tc") \
        else cell_or_para._p.get_or_add_pPr()
    target.append(el)


def set_borders(tbl, hexcolor="D8D2D4", sz=4):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), hexcolor)
        borders.append(e)
    tblPr.append(borders)


def no_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge); e.set(qn("w:val"), "none"); e.set(qn("w:sz"), "0")
        borders.append(e)
    tblPr.append(borders)


def add_toc_field(doc):
    """A real Word TOC field. Word offers to update it on open (or F9), and it
    stays live if the reader edits headings -- which a baked-in list would not."""
    p = doc.add_paragraph()
    r = p.add_run()
    for kind, text in (("begin", None), (None, 'TOC \\o "1-3" \\h \\z \\u'), ("separate", None)):
        if kind:
            fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), kind); r._r.append(fld)
        else:
            it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
            it.text = text; r._r.append(it)
    hint = p.add_run("  [In Word: right-click here and choose Update Field to build the contents.]")
    hint.italic = True; hint.font.size = Pt(8.5); hint.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end"); r._r.append(end)


INLINE_SKIP = {"script", "style", "button", "svg", "nav"}

# Block-level tags. A container that holds NONE of these is a text block and has
# to be emitted as a single paragraph -- recursing into it would turn each inline
# <strong>/<em> into its own paragraph AND silently drop the plain text between
# them, because text nodes are not Tags. Found 2026-08-20: a <div class="callout">
# reading "Objective (d) wording. The syllabus says vascular changes; the slide
# says vascular effects." came out as three orphan fragments with the sentence
# connecting them gone. That is content loss, not a layout nit.
BLOCK_TAGS = {"p", "div", "section", "article", "main", "header", "aside", "table",
              "ul", "ol", "li", "figure", "h1", "h2", "h3", "h4", "h5", "h6",
              "blockquote", "pre", "tr", "td", "th", "thead", "tbody"}


def has_block_child(node):
    return any(isinstance(c, Tag) and c.name in BLOCK_TAGS for c in node.children)


# Spans the site styles as display:block (or as a label heading). Emitted as
# plain inline runs they weld onto the text beside them -- ".deck" turned
# "Lecture 2 | Slide 50" + "2. General Dermatology I.pptx" into "Slide 502.
# General Dermatology I.pptx", and ".labs-h" produced "Labs to orderNone
# routinely." Word has no CSS, so the line break has to be made explicit.
BREAK_BEFORE = {"deck", "labs", "labs-h", "pt", "fg-name", "fg-cite", "dup", "tag"}
LABEL_CLASSES = {"labs-h", "fg-name"}
NUMBER_CLASSES = {"n", "num", "step"}
# .script is the say-this-out-loud line and .hint the coaching note; both sit
# straight after the step text and weld to it ("...by name and role\"Hello, I am").
BREAK_BEFORE = BREAK_BEFORE | {"key", "script", "hint"}
QUOTE_CLASSES = {"script"}


def emit_runs(par, node, bold=False, italic=False, hilite=False, mono=False):
    """Walk inline content, preserving bold/italic/highlight rather than
    flattening it -- the highlights are the professor-emphasis marks and carry
    meaning."""
    for child in node.children:
        if isinstance(child, NavigableString):
            txt = str(child)
            if not txt.strip() and not txt.startswith(" "):
                continue
            r = par.add_run(re.sub(r"\s+", " ", txt))
            r.bold, r.italic = bold, italic
            if mono:
                r.font.name = "Consolas"; r.font.size = Pt(9)
            if hilite:
                hl = OxmlElement("w:highlight"); hl.set(qn("w:val"), "yellow")
                r._r.get_or_add_rPr().append(hl)
            continue
        if not isinstance(child, Tag) or child.name in INLINE_SKIP:
            continue
        n = child.name
        if n == "br":
            par.add_run().add_break()
            continue
        kls = set(child.get("class") or [])
        if n == "input":
            # A run-sheet's checkboxes are the point of it. Keep them as real
            # boxes so the Word copy stays tickable.
            if child.get("type", "").lower() == "checkbox":
                r = par.add_run("\u2610 ")
                r.font.size = Pt(11)
            continue
        if kls & BREAK_BEFORE and par.runs:
            par.add_run().add_break()
        if kls & QUOTE_CLASSES:
            r = par.add_run(" ".join(child.get_text(" ", strip=True).split()))
            r.italic = True
            r.font.color.rgb = RGBColor(0x33, 0x55, 0x66)
            continue
        if kls & NUMBER_CLASSES:
            # <span class="n">1</span> butts straight against the step text --
            # "1Reviews the chief complaint". Give it a separator.
            r = par.add_run(child.get_text(" ", strip=True) + ".  ")
            r.bold = True
            continue
        emit_runs(par, child,
                  bold or bool(kls & LABEL_CLASSES) or n in ("b", "strong", "th"),
                  italic or n in ("i", "em", "cite"),
                  hilite or (n == "mark"),
                  mono or n in ("code", "kbd", "samp"))
        # A label span needs a break on BOTH sides: breaking only before it
        # still welds it to what follows -- "Labs to order" + "None routinely."
        # came out as "Labs to orderNone routinely."
        if kls & LABEL_CLASSES:
            par.add_run().add_break()


def fit(path, max_w, max_h):
    """Cap BOTH dimensions. Fixing width alone makes the grid ragged -- a portrait
    clinical photograph came out twice the height of the landscape one beside it,
    pushing its caption a row out of line."""
    try:
        with Image.open(path) as im:
            w, h = im.size
    except Exception:
        return {"width": Inches(max_w)}
    if w / h >= max_w / max_h:
        return {"width": Inches(max_w)}
    return {"height": Inches(max_h)}


def caption_runs(cp, cap):
    """A figcaption is <span class=fg-name>Name</span>Description<span class=fg-cite>Slide</span>.
    Emitted as plain runs those three collapse into one unreadable string --
    "Tinea capitisScaly grey patches...Lecture 6 Slide 9". Give each its own line."""
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(8)
    name = cap.find("span", class_="fg-name")
    cite = cap.find("span", class_="fg-cite")
    for el in (name, cite):
        if el: el.extract()
    if name:
        r = cp.add_run(" ".join(name.get_text(" ", strip=True).split()))
        r.bold = True; r.font.size = Pt(8)
        cp.add_run().add_break()
    body = " ".join(cap.get_text(" ", strip=True).split())
    if body:
        r = cp.add_run(body); r.font.size = Pt(7.5)
    if cite:
        if body: cp.add_run().add_break()
        r = cp.add_run(" ".join(cite.get_text(" ", strip=True).split()))
        r.font.size = Pt(7); r.italic = True
        r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)



class Builder:
    def __init__(self, doc, accent, imgdirs):
        self.doc, self.accent, self.imgdirs = doc, accent, imgdirs

    # -- text -------------------------------------------------------
    def para(self, node, style="Normal", size=None, space_after=6):
        p = self.doc.add_paragraph(style=style)
        emit_runs(p, node)
        if not p.runs:
            self.doc._body._body.remove(p._p); return None
        p.paragraph_format.space_after = Pt(space_after)
        if size:
            for r in p.runs: r.font.size = Pt(size)
        return p

    def heading(self, node, level):
        text = " ".join(node.get_text(" ", strip=True).split())
        h = self.doc.add_heading(level=level)
        run = h.add_run(text)
        run.font.color.rgb = RGBColor.from_string(self.accent)
        run.font.name = "Georgia"
        h.paragraph_format.space_before = Pt(14 if level <= 1 else 10)
        h.paragraph_format.space_after = Pt(5)
        return h

    # -- tables -----------------------------------------------------
    def table(self, node):
        rows = node.find_all("tr")
        if not rows: return
        ncol = max(len(r.find_all(["td", "th"])) for r in rows)
        if not ncol: return
        t = self.doc.add_table(rows=0, cols=ncol)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_borders(t)
        for ri, tr in enumerate(rows):
            cells = tr.find_all(["td", "th"])
            row = t.add_row()
            for ci in range(ncol):
                cell = row.cells[ci]
                cell.paragraphs[0].text = ""
                if ci < len(cells):
                    src = cells[ci]
                    p = cell.paragraphs[0]
                    # PICTURES INSIDE CELLS. The dermatology comparison chart is
                    # 133 photographs in a <td class="pic">, and emit_runs only
                    # emits text -- so every one of them was silently dropped and
                    # the word-coverage gate passed at 100%, because an image is
                    # not a word. Handle the picture first, then the text.
                    pics = src.find_all("img")
                    if pics:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for im in pics:
                            path = self.resolve(im.get("src") or "")
                            if not path: continue
                            try:
                                p.add_run().add_picture(path, **fit(path, 1.35, 1.15))
                            except Exception:
                                pass
                        for im in pics: im.extract()
                        p = cell.add_paragraph()
                    emit_runs(p, src)
                    p.paragraph_format.space_after = Pt(2)
                    for r in p.runs: r.font.size = Pt(9)
                    is_head = src.name == "th"
                    if is_head:
                        for r in p.runs:
                            r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        shade(cell, self.accent)
                    elif "h" in (src.get("class") or []):
                        for r in p.runs: r.bold = True
                        shade(cell, "F1EEEF")
                    # colspan
                    span = src.get("colspan")
                    if span and span.isdigit() and ci + int(span) <= ncol:
                        for k in range(1, int(span)):
                            cell.merge(row.cells[ci + k])
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # -- boxes ------------------------------------------------------
    def box(self, node, label, fill, border, label_color, drop_heading=False):
        t = self.doc.add_table(rows=1, cols=1)
        set_borders(t, border, sz=8)
        cell = t.rows[0].cells[0]
        shade(cell, fill)
        first = cell.paragraphs[0]
        lr = first.add_run(label)
        lr.bold = True; lr.font.size = Pt(9)
        lr.font.color.rgb = RGBColor.from_string(label_color)
        first.paragraph_format.space_after = Pt(4)
        inner = Builder(_CellDoc(cell), self.accent, self.imgdirs)
        if not has_block_child(node):
            # e.g. a callout that is one sentence of mixed inline markup --
            # emit it whole rather than per-child, or the prose between the
            # <strong> and <em> runs is lost.
            inner.para(node)
            self.doc.add_paragraph().paragraph_format.space_after = Pt(4)
            return
        for child in node.children:
            if not isinstance(child, Tag):
                continue
            # The io-box opens with its own "Instructional Objectives" heading,
            # which would print immediately under the label we just wrote.
            if drop_heading and child.name in ("h2", "h3", "h4") and \
               "objective" in child.get_text(" ", strip=True).lower():
                continue
            inner.dispatch(child, in_box=True)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # -- figures ----------------------------------------------------
    def figgrid(self, node):
        figs = node.find_all("figure")
        if not figs: return
        per = 3
        t = self.doc.add_table(rows=0, cols=per)
        no_borders(t)
        for i in range(0, len(figs), per):
            chunk = figs[i:i + per]
            row = t.add_row()
            for ci in range(per):
                cell = row.cells[ci]
                cell.paragraphs[0].text = ""
                if ci >= len(chunk): continue
                fig = chunk[ci]
                img = fig.find("img")
                if img and img.get("src"):
                    path = self.resolve(img["src"])
                    if path:
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        try:
                            p.add_run().add_picture(path, **fit(path, 1.85, 1.45))
                        except Exception:
                            pass
                cap = fig.find("figcaption")
                if cap:
                    caption_runs(cell.add_paragraph(), cap)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def figure(self, node):
        img = node.find("img")
        if not img or not img.get("src"): return
        path = self.resolve(img["src"])
        if not path: return
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(path, width=Inches(4.2))
        except Exception:
            return
        cap = node.find("figcaption")
        if cap:
            caption_runs(self.doc.add_paragraph(), cap)

    def resolve(self, src):
        if src.startswith("data:"): return None
        src = urllib.parse.unquote(src.split("?")[0])
        for d in self.imgdirs:
            p = os.path.join(d, src)
            if os.path.exists(p):
                return shrink(p)
        return None

    # -- lists ------------------------------------------------------
    def lst(self, node, level=0):
        style = "List Bullet" if node.name == "ul" else "List Number"
        for li in node.find_all("li", recursive=False):
            nested = [c for c in li.find_all(["ul", "ol"], recursive=False)]
            for n in nested: n.extract()
            p = self.doc.add_paragraph(style=style if level == 0 else style + " %d" % min(level + 1, 3))
            emit_runs(p, li)
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs: r.font.size = Pt(10)
            for n in nested:
                self.lst(n, level + 1)

    # -- dispatch ---------------------------------------------------
    def dispatch(self, node, in_box=False):
        if not isinstance(node, Tag) or node.name in INLINE_SKIP:
            return
        cls = node.get("class") or []
        n = node.name
        if n in ("h1",):        self.heading(node, 0 if not in_box else 3)
        elif n == "h2":         self.heading(node, 1)
        elif n == "h3":         self.heading(node, 2 if not in_box else 3)
        elif n == "h4":         self.heading(node, 3)
        elif n == "table":      self.table(node)
        elif n in ("ul", "ol"): self.lst(node)
        elif n == "figure":     self.figure(node)
        elif n == "p":
            small = "tag" in cls
            self.para(node, size=8.5 if small else None)
        elif n in ("div", "section", "article", "main", "header", "aside"):
            if "figgrid" in cls:   self.figgrid(node); return
            if "io-box" in cls:
                self.box(node, "INSTRUCTIONAL OBJECTIVES", "EEF5F4", self.accent,
                         self.accent, drop_heading=True); return
            if "prof-flag" in cls:
                self.box(node, "★  PROFESSOR EMPHASIS", "FFFBEF", "D4A017", "8A6205"); return
            if "callout" in cls or "note" in cls:
                self.box(node, "NOTE", "F5F2EC", "B9A98B", "6B5B3E"); return
            if not has_block_child(node):
                self.para(node); return          # text block -- keep it as one paragraph
            for child in node.children:
                self.dispatch(child, in_box)
        elif n in ("blockquote",):
            p = self.para(node)
            if p: p.paragraph_format.left_indent = Inches(0.3)
        elif n in ("span", "b", "strong", "em", "i", "mark", "a", "code", "small"):
            # bare inline at block level
            if node.get_text(strip=True):
                self.para(node)
        else:
            if node.name not in BLOCK_TAGS and not has_block_child(node):
                if node.get_text(strip=True): self.para(node)
                return
            for child in node.children:
                self.dispatch(child, in_box)


class _CellDoc:
    """Lets Builder write into a table cell using the same calls as a document."""
    def __init__(self, cell): self.cell = cell
    def add_paragraph(self, style=None):
        p = self.cell.add_paragraph(style=style) if style else self.cell.add_paragraph()
        return p
    def add_heading(self, level=1):
        p = self.cell.add_paragraph()
        return p
    def add_table(self, rows, cols):
        return self.cell.add_table(rows=rows, cols=cols)
    @property
    def _body(self): return self.cell


_SHRUNK = {}

def shrink(path, maxw=900):
    """Downscale for Word. A guide with 131 photographs would otherwise produce
    a document too big to email, which defeats the point of a portable copy."""
    if path in _SHRUNK: return _SHRUNK[path]
    out = os.path.join(TMP, re.sub(r"[^A-Za-z0-9._-]", "_", os.path.relpath(path, ROOT)))
    out = os.path.splitext(out)[0] + ".jpg"
    try:
        im = Image.open(path).convert("RGB")
        if im.width > maxw:
            im = im.resize((maxw, round(im.height * maxw / im.width)), Image.LANCZOS)
        os.makedirs(os.path.dirname(out), exist_ok=True)
        im.save(out, "JPEG", quality=82, optimize=True)
        _SHRUNK[path] = out
        return out
    except Exception:
        _SHRUNK[path] = path
        return path


TMP = os.path.join("/private/tmp/claude-501/-Users-jaxonluke/"
                   "8623a091-045a-42b8-8052-ca7d2eb04188/scratchpad", "docx-imgs")


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def advertise(html_path, docx_name):
    """Add <link rel="alternate" type=...docx> to the page head, idempotently."""
    src = open(html_path, encoding="utf-8").read()
    href = urllib.parse.quote(docx_name)
    tag = ('<link rel="alternate" type="%s" href="%s" title="Editable Word copy">'
           % (DOCX_MIME, href))
    if tag in src:
        return False
    src = re.sub(r'\n?\s*<link rel="alternate" type="%s"[^>]*>' % re.escape(DOCX_MIME),
                 "", src)
    m = re.search(r"</head>", src, re.I)
    assert m, "no </head> in %s" % html_path
    src = src[:m.start()] + "  " + tag + "\n" + src[m.start():]
    open(html_path, "w", encoding="utf-8").write(src)
    return True



def accent_of(soup, html):
    for pat in (r"--accent:\s*(#[0-9a-fA-F]{6})", r"navy=\"(#[0-9a-fA-F]{6})",
                r"--navy:\s*(#[0-9a-fA-F]{6})", r"\.deck-title[^}]*color:\s*(#[0-9a-fA-F]{6})"):
        m = re.search(pat, html)
        if m: return m.group(1)[1:].upper()
    return "17494B"


def convert(relpath):
    src = os.path.join(ROOT, relpath)
    html = open(src, encoding="utf-8").read()
    soup = BeautifulSoup(html, "html.parser")
    for bad in soup(["script", "style", "noscript", "svg", "button"]):
        bad.decompose()
    # Navigation is dropped: the Word TOC field replaces it, and a link rail
    # flattens into an unreadable run-on ("Sensitivity & SpecificityPretest &
    # Posttest Probability...") because the anchors carry no separators.
    # `.toc` covers both the guide's <nav class="toc"> and the cram sheet's
    # <div class="toc"> rail.
    # The OSCE run-sheets are self-contained pages that reimplement the theme
    # controls locally (see the self-contained-pages convention), so they carry
    # their own settings drawer and back link that theme.js's selectors miss.
    for sel in (".toc", "nav.toc", ".guide-back-bar", "#corner-actions", "#tts-bar",
                ".tts-bar", ".back-bar", "#back-link", ".jump", ".quicklinks",
                "#h2t-corner", "#h2t-settings", ".h2t-corner", ".sitebar"):
        for el in soup.select(sel):
            el.decompose()

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"; st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.12

    # LANDSCAPE for wide documents. The dermatology comparison chart is a
    # six-column table -- picture, condition, manifestation, testing, treatment,
    # education -- and on portrait letter every column becomes a ribbon two words
    # wide. Decided from the content, not from a filename, so any future wide
    # reference sheet gets the same treatment.
    widest = max([len(r.find_all(["td", "th"]))
                  for t in soup.find_all("table") for r in t.find_all("tr")] or [0])
    landscape = widest >= 5
    for sec in doc.sections:
        if landscape:
            sec.orientation = WD_ORIENT.LANDSCAPE
            sec.page_width, sec.page_height = sec.page_height, sec.page_width
            sec.left_margin = sec.right_margin = Inches(0.45)
            sec.top_margin = sec.bottom_margin = Inches(0.45)
        else:
            sec.left_margin = sec.right_margin = Inches(0.7)
            sec.top_margin = sec.bottom_margin = Inches(0.6)

    accent = accent_of(soup, html)
    imgdirs = [os.path.dirname(src), ROOT]
    # Count source images NOW. The builder extract()s <img> tags out of the soup
    # as it consumes them, so counting afterwards returns zero and the image gate
    # silently passes -- which is what it did on the first run, reporting
    # "133/0 images" for the chart. A gate that measures after the fact measures
    # nothing.
    want_imgs = len([i for i in soup.find_all("img")
                     if i.get("src") and not i["src"].startswith("data:")])
    src_words = words_of(soup.get_text(" ", strip=True))

    title = soup.find("h1")
    tt = doc.add_heading(level=0)
    tr = tt.add_run(" ".join(title.get_text(" ", strip=True).split()) if title else relpath)
    tr.font.color.rgb = RGBColor.from_string(accent); tr.font.name = "Georgia"

    sub = doc.add_paragraph()
    sr = sub.add_run("Editable Word copy of the PA_Quizzes study aid. "
                     "Headings use real Word styles, so the navigation pane and the contents "
                     "field below both work, and restyling one heading restyles them all.")
    sr.italic = True; sr.font.size = Pt(8.5)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    add_toc_field(doc)
    doc.add_paragraph()

    b = Builder(doc, accent, imgdirs)
    body = soup.find("main") or soup.find("body") or soup
    if title and title.parent: title.extract()
    for child in body.children:
        b.dispatch(child)

    out = os.path.splitext(src)[0] + ".docx"
    doc.save(out)
    advertise(src, os.path.basename(out))

    # CONTENT COVERAGE GATE. The real risk in an HTML->Word conversion is not an
    # ugly table, it is text that quietly disappears -- which is exactly what
    # happened to a <div class="callout"> before the block/inline rule was added.
    # So compare the words actually present in the .docx against the words in the
    # source page, and refuse to ship a lossy conversion.
    lost, cov = coverage(src_words, out)
    if cov < 0.985:
        raise SystemExit("LOSSY CONVERSION for %s: only %.1f%% of source words survived.\n"
                         "  first missing: %s" % (relpath, cov * 100, lost[:12]))

    # IMAGES NEED THEIR OWN GATE. The word check above passed at 100% on the
    # comparison chart while every one of its 133 photographs was being dropped,
    # because an image is not a word. Anything that can go missing needs to be
    # counted, not assumed.
    want = want_imgs
    from docx import Document as _D
    got = len(_D(out).inline_shapes)
    if want and got < want * 0.95:
        raise SystemExit("IMAGES LOST for %s: %d of %d embedded." % (relpath, got, want))
    return out, os.path.getsize(out), cov, got, want


def docx_words(path):
    from docx import Document
    d = Document(path)
    parts = [p.text for p in d.paragraphs]
    def walk(tbl):
        for row in tbl.rows:
            for c in row.cells:
                parts.extend(p.text for p in c.paragraphs)
                for inner in c.tables: walk(inner)
    for t in d.tables: walk(t)
    return parts


def words_of(text):
    return [w for w in re.findall(r"[A-Za-z][A-Za-z'-]{2,}", text.lower())]


def coverage(src, docx_path):
    got = set(words_of(" ".join(docx_words(docx_path))))
    from collections import Counter
    missing = [w for w in Counter(src) if w not in got]
    kept = sum(1 for w in src if w in got)
    return missing, (kept / len(src) if src else 1.0)


def main():
    docs = in_scope_documents()
    if "--list" in sys.argv:
        print("Semester 2+ classes:", ", ".join(sorted(semester_scope())))
        for slug, p in docs: print("   %-16s %s" % (slug, p))
        print("%d document(s) in scope" % len(docs))
        return
    os.makedirs(TMP, exist_ok=True)
    total = 0
    for slug, rel in docs:
        out, size, cov, got, want = convert(rel)
        total += size
        print("  %-58s %6.2f MB  %.1f%% words  %d/%d images" %
              (os.path.relpath(out, ROOT), size / 1048576, cov * 100, got, want))
    print("\n%d Word document(s), %.1f MB total" % (len(docs), total / 1048576))


if __name__ == "__main__":
    main()
